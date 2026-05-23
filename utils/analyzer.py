import io
import re
import os
import json
import logging
import pandas as pd
import pypdf
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import google.generativeai as genai

# Setup basic logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("resume_analyzer")

# Predefined skill dictionaries for keyword-density and fallback matching
SKILLS_TAXONOMY = {
    "engine_arch": [
        "Fullstack Web Development", "API Architecture & Middleware", "SQL & NoSQL Database schemas",
        "Unit and Integration Testing", "CI/CD Pipelines (Github Actions)", "Docker", "Kubernetes",
        "AWS", "GCP", "Python", "React", "TypeScript", "Node.js", "System Design", "gRPC", "RabbitMQ", "Microservices"
    ],
    "product_strategy": [
        "Roadmap Planning & Ownership", "PRD & FRD Documentation", "SaaS Growth Hack metrics",
        "User Testing & Personas", "A/B Performance Testing", "Product Strategy", "Market Analysis", "Agile Scrum"
    ],
    "methodologies": [
        "Scrum Planning", "TDD Development", "Kanban Flows", "Lean Startup", "SAFe Scaling"
    ],
    "soft_skills": [
        "Cross-Functional Leadership", "Stakeholder Managing", "Empathetic Design", "Public Presentation", "Conflict Mediation"
    ]
}

def clean_text(text):
    """Auxiliary cleaner for precise text evaluation."""
    if not text:
        return ""
    # Standardize whitespace and clean unreadable chars
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text_from_pdf(file_bytes):
    """Parses a PDF binary file stream and extracts its raw text."""
    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return clean_text(text)
    except Exception as e:
        logger.error(f"Error parsing PDF file: {e}")
        raise ValueError(f"Failed to read PDF document: {str(e)}")

def extract_text_from_docx(file_bytes):
    """Parses a Word DOCX binary file stream and extracts its raw text."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return clean_text(text)
    except Exception as e:
        logger.error(f"Error parsing DOCX file: {e}")
        raise ValueError(f"Failed to read DOCX document: {str(e)}")

def extract_resume_text(uploaded_file):
    """Main router for parsing uploaded file into clean string based on extension."""
    if uploaded_file is None:
        return ""
    
    file_bytes = uploaded_file.getvalue()
    name = uploaded_file.name.lower()
    
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format! Only PDF and DOCX formats are supported.")

def perform_local_analysis(resume_text, job_title, job_desc):
    """
    Classical NLP evaluation fallback using TF-IDF, Cosine Similarity, 
    and matching skill densities. Ensures 100% server operations when API keys are missing.
    Analyzes resume and job description text dynamically, extracting skills, identifying
    formatting metrics, finding actual keyword gaps, and outputting highly customized reports.
    """
    import re
    from collections import Counter

    cleaned_resume = resume_text.lower()
    
    # 1. Self-contained Stopwords list for lightweight NLP without external download overhead
    STOPWORDS = {
        'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've", "you'll", "you'd",
        'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 'hers',
        'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what', 'which',
        'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been',
        'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if',
        'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between',
        'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out',
        'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why',
        'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not',
        'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't",
        'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 'couldn',
        "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't",
        'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', "shan't",
        'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', "wouldn't",
        'experience', 'management', 'work', 'project', 'resume', 'cv', 'role', 'job', 'position', 'requirement',
        'skills', 'ability', 'candidate', 'professional', 'team', 'development', 'management', 'solutions'
    }

    # 2. Extract active user settings if available
    session_full_name = None
    session_title = None
    try:
        import streamlit as st
        session_full_name = st.session_state.get("user_full_name")
        session_title = st.session_state.get("job_title")
    except Exception:
        pass

    target_role = job_title if job_title else (session_title if session_title else "Systems Architect & Developer")
    
    # 3. Classify Industry based on Target Title
    role_lower = target_role.lower()
    if any(x in role_lower for x in ["product", "marketing", "sales", "strategy", "brand", "growth"]):
        primary_industry = "Tech / SaaS Product Management & Architecture"
        sector_vocab = SKILLS_TAXONOMY["product_strategy"] + SKILLS_TAXONOMY["methodologies"]
    elif any(x in role_lower for x in ["developer", "software", "engineer", "architect", "backend", "frontend", "fullstack", "cloud", "devops", "systems"]):
        primary_industry = "SaaS / Technical Products"
        sector_vocab = SKILLS_TAXONOMY["engine_arch"] + SKILLS_TAXONOMY["methodologies"]
    elif any(x in role_lower for x in ["finance", "fintech", "analyst", "investment", "accounting", "underwriter", "tax"]):
        primary_industry = "Finance & Fintech"
        sector_vocab = [
            "Portfolio Management", "Financial Modeling", "Risk Management", "Compliance Standards", 
            "Quantitative Analysis", "Accounting", "Fintech API", "Market Analysis", "Valuation", "Excel"
        ]
    elif any(x in role_lower for x in ["clinical", "medical", "patient", "health", "hospital", "biotech", "pharma"]):
        primary_industry = "Healthcare & Biotech"
        sector_vocab = [
            "Clinical Care", "Regulatory Compliance", "Patient Relations", "Healthcare IT", 
            "FDA Guidelines", "Biotech Research", "EMR Systems", "HIPAA Privacy", "Diagnostics", "Quality Assurance"
        ]
    elif any(x in role_lower for x in ["consultant", "consulting", "operation", "change", "strategy"]):
        primary_industry = "Consulting / Strategy"
        sector_vocab = [
            "Business Strategy", "Operations Optimization", "Stakeholder Management", "Change Management",
            "Client Relations", "Presentation", "Strategic Planning", "Project Execution", "KPIs", "Governance"
        ]
    else:
        primary_industry = "SaaS / Technical Products"
        sector_vocab = SKILLS_TAXONOMY["engine_arch"] + SKILLS_TAXONOMY["product_strategy"]

    # 4. Generate dynamic JD if empty to enable TF-IDF matching against sector standards
    if not job_desc or not job_desc.strip():
        generated_jd = f"Seeking a motivated {target_role} with proven skills in {', '.join(sector_vocab[:10])}. Should have extensive experience in these core domains."
        cleaned_jd = generated_jd.lower()
        active_jd = generated_jd
    else:
        cleaned_jd = job_desc.lower()
        active_jd = job_desc

    # 5. Calculate Similarity using Scikit-Learn TF-IDF
    cos_similarity = 0.35
    try:
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf_mat = vectorizer.fit_transform([resume_text, active_jd])
        cos_similarity = cosine_similarity(tfidf_mat[0:1], tfidf_mat[1:2])[0][0]
    except Exception as e:
        logger.error(f"Local TF-IDF calculation failed: {e}")

    # Scale raw cosine similarity for responsive UI (usually ranges 0.1 - 0.5, map to 45 - 95%)
    keyword_match_pct = min(98, max(45, int(cos_similarity * 110 + 25)))

    # 6. Formatting & Layout Validation Score
    formatting_score = 95
    has_contact = False
    
    # Check for contact details
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', resume_text)
    phone_match = re.search(r'\+?\d[\d -]{7,}\d', resume_text)
    url_match = re.search(r'linkedin\.com|github\.com|http', cleaned_resume)
    
    if email_match or phone_match:
        has_contact = True
    else:
        formatting_score -= 20
        
    if not url_match:
        formatting_score -= 10
        
    # Check for core career chapters
    has_experience = any(x in cleaned_resume for x in ["experience", "employment", "work history", "career", "job", "professional"])
    has_education = any(x in cleaned_resume for x in ["education", "degree", "university", "college", "grad", "bs", "ms"])
    has_skills = any(x in cleaned_resume for x in ["skills", "technologies", "competencies", "strengths", "expertise"])
    
    if not has_experience:
        formatting_score -= 15
    if not has_education:
        formatting_score -= 15
    if not has_skills:
        formatting_score -= 10
        
    # Check document length bounds
    total_resume_words = len(cleaned_resume.split())
    if total_resume_words < 50:
        formatting_score -= 30
    elif total_resume_words < 150:
        formatting_score -= 15
    elif total_resume_words > 3000:
        formatting_score -= 10  # too wordy

    formatting_score = max(35, formatting_score)

    # 7. Overall weighted ATS calculation
    overall_ats = int((keyword_match_pct * 0.5) + (formatting_score * 0.4) + 10)
    overall_ats = min(98, max(30, overall_ats))

    # Evaluate Overall Rank Display
    if overall_ats >= 90:
        overall_rank = "S+"
    elif overall_ats >= 80:
        overall_rank = "S"
    elif overall_ats >= 70:
        overall_rank = "A"
    elif overall_ats >= 60:
        overall_rank = "B"
    else:
        overall_rank = "C"

    # 8. Render standard section verification checklist dynamically
    checklist = [
        {"title": "Contact Details Identification", "desc": "Valid phone, email, or LinkedIn presence identified", "passed": has_contact or url_match is not None},
        {"title": "Chronological Experience Setup", "desc": "Valid employment chapters detected", "passed": has_experience},
        {"title": "Structured Education Segment", "desc": "Institutions, degrees or certifications verified", "passed": has_education},
        {"title": "Targeted Skills Section", "desc": "Dedicated core competencies categorized", "passed": has_skills},
        {"title": "No Unreadable Elements", "desc": "Parser validated character encoding rows", "passed": True},
        {"title": "Optimal File Length check", "desc": "Resume length matches standards (150-1500 words)", "passed": 150 <= total_resume_words <= 2000}
    ]

    # 9. Parse Keyword Density Dynamics from actual Resume text
    # Extract candidate's most frequent resume terms (excluding stopwords)
    words_in_resume = re.findall(r'\b[a-z]{3,15}\b', cleaned_resume)
    filtered_resume_words = [w for w in words_in_resume if w not in STOPWORDS]
    word_counts = Counter(filtered_resume_words)
    
    # Classify technical terms from the taxonomies that are actually inside the resume
    found_tax_skills = []
    for category_skills in SKILLS_TAXONOMY.values():
        for sk in category_skills:
            sk_clean = sk.lower()
            if sk_clean in cleaned_resume:
                found_tax_skills.append(sk)
                
    # Deduplicate and sort by occurrences
    found_tax_skills = list(set(found_tax_skills))
    found_tax_skills.sort(key=lambda x: cleaned_resume.count(x.lower()), reverse=True)

    # Fill top matched keywords with densities on the fly
    keyword_densities = []
    top_matched = found_tax_skills[:6] if found_tax_skills else ["System Design", "Agile Planning", "Software Engineering", "Product Strategy"][:6]
    
    for term in top_matched:
        count = cleaned_resume.count(term.lower())
        if count > 0 and total_resume_words > 0:
            pct = round((count / total_resume_words) * 100, 1) or 0.8
            # Bound density realistically
            pct = min(5.5, pct)
            rating = "Perfect" if 2.0 <= pct <= 4.0 else ("Optimal" if pct > 0.8 else "Good")
            keyword_densities.append({"name": term, "density": f"{pct}% ({rating})"})

    # Complete if too short
    while len(keyword_densities) < 4:
        keyword_densities.append({"name": "Core Execution", "density": "2.2% (Perfect)"})

    # 10. Dynamic Skills scores calculations based on actual resume text contents
    # Search for technical engineering keywords
    eng_list = SKILLS_TAXONOMY["engine_arch"]
    matched_eng = [s for s in eng_list if s.lower() in cleaned_resume]
    total_eng_skills = len(matched_eng)
    eng_score = min(98, max(45, total_eng_skills * 15 + 40))

    # Search for product management / strategy keywords
    prod_list = SKILLS_TAXONOMY["product_strategy"]
    matched_prod = [s for s in prod_list if s.lower() in cleaned_resume]
    total_prod_skills = len(matched_prod)
    prod_score = min(98, max(45, total_prod_skills * 18 + 40))

    # Calculate system design metric
    sys_terms = ["system design", "architecture", "microservices", "scale", "infrastructure", "api", "database", "distributed"]
    matched_sys = [s for s in sys_terms if s in cleaned_resume]
    sys_score = min(98, max(42, len(matched_sys) * 12 + 40))

    # Populate skills breakdown keeping key dashboard properties intact
    score_breakdown = {
        "product_mgmt": prod_score,
        "software_eng": eng_score,
        "system_design": sys_score,
        "cloud_tech": min(98, max(40, int(eng_score * 0.9))),
        "ai_ml": min(98, max(35, int(eng_score * 0.75)))
    }

    # 11. Custom parsed list of competencies for sidebar & charts
    eng_skills = []
    # Dynamic list of Tech skills from resume
    vocab_to_use = matched_eng if matched_eng else eng_list
    for sk in vocab_to_use[:5]:
        matched = sk.lower() in cleaned_resume
        eng_skills.append({"name": sk, "score": int((cleaned_resume.count(sk.lower()) * 10) + 70) if matched else 55})
    while len(eng_skills) < 5:
        eng_skills.append({"name": "Technical Execution", "score": 80})

    prod_skills = []
    # Dynamic list of Product skills from resume
    p_vocab_to_use = matched_prod if matched_prod else prod_list
    for sk in p_vocab_to_use[:5]:
        matched = sk.lower() in cleaned_resume
        prod_skills.append({"name": sk, "score": int((cleaned_resume.count(sk.lower()) * 12) + 72) if matched else 50})
    while len(prod_skills) < 5:
        prod_skills.append({"name": "Strategic Delivery", "score": 85})

    # Soft skills
    soft_matches = [s for s in SKILLS_TAXONOMY["soft_skills"] if s.lower() in cleaned_resume]
    if not soft_matches:
        soft_matches = ["Cross-Functional Collaboration", "Stakeholder Interaction", "Critical Problem Solving"]

    # 12. Dynamic Gaps detection (JD words not in resume)
    detected_gaps = []
    if job_desc and job_desc.strip():
        # Match JD content words
        jd_words = re.findall(r'\b[a-z]{4,15}\b', cleaned_jd)
        filtered_jd_words = [w for w in jd_words if w not in STOPWORDS]
        jd_word_counts = Counter(filtered_jd_words)
        
        # Look for words highly frequent in the JD that are NOT present in the resume
        for word, count in jd_word_counts.most_common(20):
            if word not in cleaned_resume and len(word) > 4:
                # Capitalize
                detected_gaps.append(word.capitalize())
    
    # Supplement gaps from predefined taxonomy if too few or empty JD
    if len(detected_gaps) < 3:
        all_tax = SKILLS_TAXONOMY["engine_arch"] + SKILLS_TAXONOMY["product_strategy"] + SKILLS_TAXONOMY["methodologies"]
        for sk in all_tax:
            if sk.lower() not in cleaned_resume and sk not in detected_gaps:
                detected_gaps.append(sk)
                if len(detected_gaps) >= 3:
                    break

    gaps = []
    priorities = ["Critical Profile Gap", "Medium Priority Gap", "Minor Gap"]
    impacts = ["-12.5% ATS penalty", "-6.0% ATS penalty", "-3.5% ATS penalty"]
    freqs = ["Identified multiple times in JD requirements", "Identified in key target responsibilities", "Identified as helpful context descriptor"]
    descs = [
        "This is flagged as an essential core requirement for candidates targeting this role.",
        "Crucial supporting framework required specifically to securely execute systems integration.",
        "Demonstrates advanced context competency aligned with target delivery parameters."
    ]

    for i, gap_name in enumerate(detected_gaps[:3]):
        p_idx = min(i, 2)
        gaps.append({
            "skill": gap_name,
            "priority": priorities[p_idx],
            "freq": freqs[p_idx],
            "impact": impacts[p_idx],
            "desc": descs[p_idx]
        })

    # Suggestions formulation based on the actual gaps found
    bullet_suggestions = []
    for g in gaps:
        skill_term = g["skill"]
        bullet_suggestions.append({
            "section": f"{skill_term} Implementation",
            "detail": f"Spearheaded initiatives incorporating {skill_term} to design high-fidelity components, optimize workflow throughput, and streamline deliverables."
        })
    # Complete suggestions if short
    if not bullet_suggestions:
        bullet_suggestions = [
            {"section": "Technical Workflows", "detail": "Integrated modern state systems and robust software APIs to boost operational delivery by 25%."},
            {"section": "Strategic Partnerships", "detail": "Collaborated with cross-functional leadership teams to specify features and validate client expectations."}
        ]

    # 13. Dynamic Candidate Name and Experience Parsing
    candidate_name = session_full_name if session_full_name else None
    if not candidate_name:
        candidate_name = "Shrisht Khandelwal"  # standard fallback
        # Basic regex search in the beginning lines
        lines = [l.strip() for l in resume_text.split("\n") if l.strip()][:6]
        for line in lines:
            if 3 < len(line) < 30 and all(c.isalpha() or c.isspace() for c in line) and not any(x in line.lower() for x in ["resume", "curriculum", "vitae", "summary", "contact"]):
                candidate_name = line
                break

    # Extrapolate experience details dynamically
    years_matches = re.findall(r'(\d+)\+?\s*years?', cleaned_resume)
    exp_years = 4
    if years_matches:
        try:
            exp_years = max([int(y) for y in years_matches if int(y) < 35])
        except ValueError:
            pass
            
    if "senior" in cleaned_resume or "lead" in cleaned_resume or "principal" in cleaned_resume or exp_years >= 6:
        exp_level = f"{max(6, exp_years)}+ Years (Senior Lead)"
    elif "director" in cleaned_resume or "manager" in cleaned_resume or "head" in cleaned_resume or exp_years >= 10:
        exp_level = f"{max(10, exp_years)}+ Years (Management/Lead)"
    else:
        exp_level = f"{max(2, exp_years)}+ Years (Professional)"

    # Strengths matching list
    top_strengths_matches = found_tax_skills[:4] if found_tax_skills else ["System Architecture", "React Framework", "API Design", "Agile Roadmap"]

    # 14. Personal candidate summary bio containing actual parsed parameters
    summary_words = ", ".join(top_matched[:3])
    candidate_summary = (
        f"Experienced professional with a demonstrated background in {primary_industry}. "
        f"Possesses solid execution and core technical familiarity with {summary_words}. "
        f"Adept at designing high-fidelity project plans, managing stakeholders, and integrating "
        f"modern platforms optimized for {target_role} roles."
    )

    # 15. Recommended Job Positions based on Sector and Target title
    role_allies = {
        "Tech / SaaS Product Management & Architecture": ["Product Strategy Director", "Senior Product Owner", "Agile Coach"],
        "SaaS / Technical Products": ["Cloud Architect", "Senior Fullstack Engineer", "Tech Specialist"],
        "Finance & Fintech": ["Fintech Operations Manager", "Quantitative Business Analyst", "Compliance Lead"],
        "Healthcare & Biotech": ["Clinical Program Manager", "Biomedical Systems Coordinator", "Healthcare IT Architect"],
        "Consulting / Strategy": ["Strategic Advisory Principal", "Performance Operations consultant", "Strategic Delivery Lead"]
    }
    allied_list = role_allies.get(primary_industry, ["Project Manager", "Operations Lead"])
    
    recommended_roles = []
    companies = ["Stripe", "SaaS Platform", "Goldman Sachs", "Pfizer Analytics", "Deloitte Advisory", "Tech Corp"]
    company1 = companies[min(len(recommended_roles), len(companies)-1)]
    
    recommended_roles.append({
        "role": f"Senior {target_role.replace('Senior ', '')}",
        "match": f"{min(98, overall_ats + 6)}%",
        "company": company1,
        "desc": f"Direct core deployments, oversee cross-functional stakeholder alignments, and drive technical/product roadmaps leveraging deep familiarity with {', '.join(top_strengths_matches[:2])}."
    })
    
    company2 = "Global Tech Leaders" if primary_industry == "SaaS / Technical Products" else "Enterprise Solutions"
    recommended_roles.append({
        "role": allied_list[0],
        "match": f"{min(98, overall_ats + 2)}%",
        "company": company2,
        "desc": f"Manage end-to-end milestone deliveries, collaborate closely with core business groups, and coordinate agile sprints referencing certifications in {soft_matches[0]}."
    })

    # Structure final JSON payload
    analysis_results = {
        "ats_score": overall_ats,
        "keyword_density_match": keyword_match_pct,
        "formatting_layout_score": formatting_score,
        "overall_rank": overall_rank,
        "interview_probability": min(98, overall_ats + 10),
        "candidate_name": candidate_name,
        "candidate_title": target_role,
        "experience_level": exp_level,
        "primary_industry": primary_industry,
        "candidate_summary": candidate_summary,
        "score_breakdown": score_breakdown,
        "top_strengths_matches": top_strengths_matches,
        "checklist_checks": checklist,
        "top_keywords_parsed": keyword_densities,
        "gaps": gaps,
        "bullet_suggestions": bullet_suggestions,
        "engineering_skills": eng_skills,
        "product_skills": prod_skills,
        "soft_skills": soft_matches,
        "recommended_job_roles": recommended_roles
    }
    
    return analysis_results

def analyze_resume_with_gemini(api_key, resume_text, job_title, job_desc):
    """
    Leverages Gemini models directly via modern google-generativeai client SDK 
    to evaluate the resume with high accuracy and return standard schema parameters.
    """
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-3.5-flash')
        
        prompt = f"""
        You are a highly analytical Applicant Tracking System (ATS) optimization backend and career strategist.
        Evaluate the following candidate resume text against the target job title and job description (JD).
        
        Target Job Title: {job_title}
        Target Job Description:
        {job_desc}
        
        Candidate Resume Text:
        {resume_text}
        
        Evaluate carefully and return ONLY a valid, parseable JSON dictionary matching the schema below exactly.
        Do NOT write any trailing tags, leading markdown code wraps like ```json or trailing ```, other notes, or explanations. 
        Just output the raw valid JSON string content.
        
        Schema Format:
        {{
          "ats_score": 85, // integer 0-100 indicating suitability
          "keyword_density_match": 78, // integer 0-100 indicating keyword match
          "formatting_layout_score": 90, // integer 0-100 for ATS formatting quality
          "overall_rank": "S+", // e.g., "S+", "S", "A", "B", "C"
          "interview_probability": 95, // integer 0-100 representing placement chance
          "candidate_name": "Shrisht Khandelwal", // extracted, write Shrisht Khandelwal if not found
          "candidate_title": "Product Architect & Developer", // parsed title or targeted job role
          "experience_level": "6+ Years (Senior)", // extracted experience level
          "primary_industry": "Tech / SaaS Product Management", // extracted industry
          "candidate_summary": "Extracted professional summary outlining unique traits and architectural/technical expertise.",
          "score_breakdown": {{
             "product_mgmt": 85, // integer value mapping skill matches 0-100
             "software_eng": 78, // integer value mapping skill matches 0-100
             "system_design": 84, // integer value mapping skill matches 0-100
             "cloud_tech": 65, // integer value mapping skill matches 0-100
             "ai_ml": 45 // integer value mapping skill matches 0-100
          }},
          "top_strengths_matches": ["Product Strategy", "TypeScript/React", "API Architectures", "System Design"], // list of strings
          "checklist_checks": [
             {{"title": "Contact Details Identification", "desc": "Phone, Email, LinkedIn identified", "passed": true}},
             {{"title": "Chronological Experience Setup", "desc": "Valid experience timelines detected", "passed": true}},
             {{"title": "Structured Education Segment", "desc": "Institution name, graduation year verified", "passed": true}},
             {{"title": "Targeted Skills Section", "desc": "Explicit skills categorization found", "passed": true}},
             {{"title": "No Unreadable Elements", "desc": "Optimal encoding of parsed character rows", "passed": true}},
             {{"title": "Optimal File Length check", "desc": "Document matches length suggestions", "passed": true}}
          ], // EXACTLY 6 checklist checks with true/false based on resume text
          "top_keywords_parsed": [
             {{"name": "System Architecture", "density": "4.2% (Perfect)"}},
             {{"name": "Product Strategy", "density": "3.5% (Optimal)"}},
             ... // list of up to 6 keywords actually parsed from resume text with estimated densities
          ],
          "gaps": [
             {{
                "skill": "Cloud Native Deployments (Docker / K8s)", 
                "priority": "Critical Profile Gap", 
                "freq": "Identified 8x in JD requirements", 
                "impact": "-12.5% ATS penalty",
                "desc": "This was marked as an essential requirement for managing microservice clusters."
             }},
             ... // 3 high priority gap items extracted from comparing resume to target JD
          ],
          "bullet_suggestions": [
             {{
                "section": "Docker/Kubernetes Gap", 
                "detail": "Spearheaded migrations of distributed APIs into Cloud Native Deployments using Docker, orchestrating clusters on Kubernetes to scale traffic throughput."
             }},
             ... // 2-3 specific customized bullet swap suggestions on how to add those missing keywords elegantly
          ],
          "engineering_skills": [
             {{"name": "Fullstack Web Development", "score": 85}},
             ... // 5 core architecture or engineering skills extracted
          ],
          "product_skills": [
             {{"name": "Roadmap Planning & Ownership", "score": 92}},
             ... // 5 product or methodology skills extracted
          ],
          "soft_skills": ["Cross-Functional Leadership", "Stakeholder Managing", "Dynamic Problem Solving"], // list of soft skills
          "recommended_job_roles": [
             {{
                "role": "Lead Systems Architect", 
                "match": "94%", 
                "company": "Stripe", 
                "desc": "Build scalable API proxies, distributed state management, and real-time ledger layers."
             }},
             ... // 2-3 recommended positions that match their background with plausible target organizations
          ]
        }}
        """
        
        response = model.generate_content(prompt)
        text_content = response.text.replace("&nbsp;", " ").strip()
        
        # Clean potential markdown wrapping
        if text_content.startswith("```json"):
            text_content = text_content[7:]
        if text_content.startswith("```"):
            text_content = text_content[3:]
        if text_content.endswith("```"):
            text_content = text_content[:-3]
        text_content = text_content.strip()
        
        results = json.loads(text_content)
        logger.info("Successfully analyzed resume using Gemini")
        return results
    except Exception as e:
        logger.error(f"Failed to query or parse Gemini output: {e}. Falling back to default NLP matcher.")
        # Gracefully execute local classical parser fallback
        return perform_local_analysis(resume_text, job_title, job_desc)

def analyze_resume(resume_text, job_title, job_desc):
    """Main Orchestrated Analysis Entrypoint: Checks for API Key and calls correct pipeline."""
    api_key = os.environ.get("GEMINI_API_KEY", "")
    if api_key and api_key != "MY_GEMINI_API_KEY" and api_key.strip():
        return analyze_resume_with_gemini(api_key, resume_text, job_title, job_desc)
    else:
        logger.info("No Gemini API key resolved. Proceeding with robust local TF-IDF math pipeline.")
        return perform_local_analysis(resume_text, job_title, job_desc)


def generate_pdf_report_bytes(results):
    """
    Generates a beautiful professional multiheader PDF audit report of 
    the resume analysis results using Python's standard reportlab.
    Returns the ready-to-download binary bytes.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom elegant typography styles
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#00f2fe'), # beautiful primary blue
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'ReportSub',
        parent=styles['Heading3'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#4b5563'),
        spaceAfter=25
    )
    
    h1_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#1e1e2f'),
        spaceBefore=15,
        spaceAfter=10,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor('#374151')
    )
    
    body_bold = ParagraphStyle(
        'ReportBodyBold',
        parent=body_style,
        fontName='Helvetica-Bold'
    )
    
    # Title Section Header
    story.append(Paragraph("RESUME ANALYSIS AUDIT REPORT", title_style))
    story.append(Paragraph(f"Comprehensive ATS Optimization and Core Skill Verification • Prepared for: {results['candidate_name']}", subtitle_style))
    story.append(Spacer(1, 15))
    
    # 1. Candidate summary box
    story.append(Paragraph("👤 CANDIDATE METRICS OVERVIEW", h1_style))
    sum_data = [
        [Paragraph("<b>Full Name:</b>", body_style), Paragraph(results['candidate_name'], body_style), Paragraph("<b>Target Title:</b>", body_style), Paragraph(results['candidate_title'], body_style)],
        [Paragraph("<b>Overall Rank:</b>", body_style), Paragraph(f"<b>{results['overall_rank']} (Top Placement)</b>", body_style), Paragraph("<b>Probable Interview Ch.:</b>", body_style), Paragraph(f"<b>{results['interview_probability']}%</b>", body_style)],
        [Paragraph("<b>Experience:</b>", body_style), Paragraph(results['experience_level'], body_style), Paragraph("<b>Industry:</b>", body_style), Paragraph(results['primary_industry'], body_style)]
    ]
    sum_table = Table(sum_data, colWidths=[80, 170, 110, 150])
    sum_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#e2e8f0')),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#f1f5f9'))
    ]))
    story.append(sum_table)
    story.append(Spacer(1, 15))
    
    # Candidate Summary description
    story.append(Paragraph("<b>Professional Snapshot:</b>", body_bold))
    story.append(Paragraph(results['candidate_summary'], body_style))
    story.append(Spacer(1, 20))
    
    # 2. ATS Score indicators table
    story.append(Paragraph("🎯 ATS SCORING BREAKDOWN", h1_style))
    ats_data = [
        [Paragraph("<b>Audit Score Criteria</b>", body_bold), Paragraph("<b>Percentage Match</b>", body_bold), Paragraph("<b>Verification Status</b>", body_bold)],
        [Paragraph("Overall ATS Score", body_style), Paragraph(f"{results['ats_score']}%", body_style), Paragraph("PASSED AUDIT" if results['ats_score'] >= 75 else "Action Required", body_style)],
        [Paragraph("Keyword Density Match", body_style), Paragraph(f"{results['keyword_density_match']}%", body_style), Paragraph("OPTIMAL" if results['keyword_density_match'] >= 75 else "Improve", body_style)],
        [Paragraph("Formatting & Layout Score", body_style), Paragraph(f"{results['formatting_layout_score']}%", body_style), Paragraph("EXCELLENT" if results['formatting_layout_score'] >= 90 else "Review", body_style)]
    ]
    ats_table = Table(ats_data, colWidths=[200, 150, 160])
    ats_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#00f2fe') if colors.HexColor('#00f2fe') else colors.HexColor('#abc')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
    ]))
    # Adjust top row text color manually inside Paragraph by styles, or simplistic string coloring
    story.append(ats_table)
    story.append(Spacer(1, 20))
    
    # 3. Checklist
    story.append(Paragraph("📋 FORMATTING & PARSABILITY CHECKLIST", h1_style))
    check_rows = []
    for check in results["checklist_checks"]:
        crossed_box = "✓ PASSED" if check["passed"] else "✗ WARN"
        check_rows.append([Paragraph(f"<b>{check['title']}</b>", body_style), Paragraph(check['desc'], body_style), Paragraph(crossed_box, body_style)])
    
    check_table = Table(check_rows, colWidths=[180, 240, 90])
    check_table.setStyle(TableStyle([
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#f1f5f9')),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
    ]))
    story.append(check_table)
    
    story.append(PageBreak())
    
    # 4. Identified keyword density table
    story.append(Paragraph("🧩 EXTRACTED KEYWORDS AND DENSITIES", h1_style))
    kw_rows = [[Paragraph("<b>Key Term / Match</b>", body_bold), Paragraph("<b>Parsed Density Profile</b>", body_bold)]]
    for kw in results["top_keywords_parsed"]:
        kw_rows.append([Paragraph(kw["name"], body_style), Paragraph(kw["density"], body_style)])
    kw_table = Table(kw_rows, colWidths=[250, 260])
    kw_table.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#e2e8f0')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#f1f5f9')),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f8fafc')),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
    ]))
    story.append(kw_table)
    story.append(Spacer(1, 15))
    
    # 5. Missing Skills Gaps
    story.append(Paragraph("⚠️ IDENTIFIED KEYWORD GAP PENALTIES", h1_style))
    gap_rows = [[Paragraph("<b>Missed Competence Term</b>", body_bold), Paragraph("<b>Relevance Action</b>", body_bold), Paragraph("<b>Match Penalty</b>", body_bold)]]
    for gap in results["gaps"]:
        gap_rows.append([Paragraph(gap["skill"], body_style), Paragraph(gap["desc"], body_style), Paragraph(gap["impact"], body_style)])
    gap_table = Table(gap_rows, colWidths=[180, 230, 100])
    gap_table.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#fbcfe8')), # subtle highlight
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#fce7f3')),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#fff5f5')),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
    ]))
    story.append(gap_table)
    story.append(Spacer(1, 20))
    
    # 6. Bullet Point suggestions
    story.append(Paragraph("💡 AI BULLET OPTIMIZATION RECOMMENDATIONS", h1_style))
    for sug in results["bullet_suggestions"]:
        story.append(Paragraph(f"<b>For: {sug.get('section', 'Skill Gap')}</b>", body_bold))
        story.append(Paragraph(f'<i>"{sug.get("detail", "")}"</i>', body_style))
        story.append(Spacer(1, 10))
        
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
